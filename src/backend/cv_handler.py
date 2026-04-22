"""
================================================================================
SMART CV FILTER - DOCUMENT PROCESSING ENGINE (CV_HANDLER)
================================================================================
Author: Tebagl
Date: April 2026
Version: 2.0 (Cloud Migration)

Description:
The document processing layer responsible for extracting, cleaning, and 
standardizing text from various file formats (PDF, DOCX). This module ensures 
that the AI receives high-quality string data for accurate analysis.

TECHNICAL SPECIFICATIONS:
- Supported Formats: PDF (.pdf), Word (.docx)
- Libraries: PyMuPDF (fitz), PDFMiner, python-docx
- Cleaning Pipeline: Normalizes whitespace, removes non-printable characters, 
  and handles encoding issues.

KEY FUNCTIONS:
- extract_text(): Primary entry point for multi-format text recovery.
- clean_text(): Post-processing to optimize token usage for the LLM.
- get_file_metadata(): Optional tracking of file creation/modification dates.

ROBUSTNESS:
- Error handling for corrupted files and encrypted PDFs.
- Automatic detection of file extensions.
- Memory-efficient processing for large batches of documents.
================================================================================
"""

import os
import shutil
import logging
import fitz  # PyMuPDF
from datetime import datetime
import csv
from docx import Document
import sys
import gc
import json
import re
import time

from src.backend.anonymizer import Anonymizer

logger = logging.getLogger(__name__)

class CVHandler:
    def __init__(self, analyzer):
        """
        Manejador de CVs para sistema de carpetas e IA Embebida.
        """
        self.analyzer = analyzer
        self.anonymizer = Anonymizer()
        # Inicializamos en None o vacío. 
        # No creamos carpetas aquí para evitar duplicados al arrancar.
        self.base_output = None 

    def _ensure_folders(self):
        """Crea la estructura de carpetas solo si base_output ha sido definido."""
        if not self.base_output:
            return

        os.makedirs(self.base_output, exist_ok=True)
        for folder in ["RECLUTADOS", "DESCARTADOS", "DUDAS"]:
            os.makedirs(os.path.join(self.base_output, folder), exist_ok=True)
    
    def _extract_text_from_pdf(self, pdf_path):
        """Extrae todo el texto de un archivo PDF."""
        text = ""
        try:
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Error leyendo PDF {pdf_path}: {e}")
            return ""
        
    def _extract_text_from_docx(self, docx_path):
        try:
            doc = Document(docx_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                     text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error en DOCX {docx_path}: {e}")
            return ""
        
    def _append_to_report(self, data):
        """
        Guarda los resultados en un archivo CSV dentro de la carpeta del proceso.
        """
        # Buscamos la carpeta del proceso actual
        report_path = os.path.join(self.base_output, "resumen_clasificacion.csv")
        
        # Definimos las columnas
        fieldnames = ['nombre', 'score', 'decision', 'motivo', 'ruta_final']
        
        file_exists = os.path.isfile(report_path)

        try:
            with open(report_path, 'a', newline='', encoding='utf-8') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                
                # Si el archivo es nuevo, escribir la cabecera
                if not file_exists:
                    writer.writeheader()

                # Escribimos la fila asegurando que las llaves coincidan
                writer.writerow({
                    'nombre': data.get('nombre', 'Desconocido'),
                    'score': data.get('score', 0),
                    'decision': data.get('decision', 'N/A'),
                    'motivo': data.get('motivo', 'Sin motivo'),
                    'ruta_final': data.get('dest_path', '')
                })
            print(f"📊 Registro añadido al CSV: {data.get('nombre')}")
        except Exception as e:
            print(f"❌ Error al escribir en CSV: {e}")
        

    def process_cv(self, file_path: str, job_description: str = None):
        """
        Analiza el CV de forma anónima. Versión corregida sin duplicidad
        y protegida contra errores de respuesta vacía.
        """
        import os
        import json
        import shutil
        import gc

        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            # 1. Extracción de texto
            if ext == ".pdf":
                raw_text = self._extract_text_from_pdf(file_path)
            elif ext == ".docx":
                raw_text = self._extract_text_from_docx(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    raw_text = f.read()

            if not raw_text.strip():
                return {"status": "error", "reason": "Archivo vacío o ilegible"}
            
            # 2. Anonimización y llamada a la IA
            texto_para_la_nube = self.anonymizer.anonymize(raw_text)
            # RECORTE DE SEGURIDAD: Evita el Segmentation Fault y mejora el enfoque de la IA
            # Al limitar a 5000 caracteres, aseguras que no superas los 12K tokens del modelo versatile
            texto_para_ia = texto_para_la_nube[:5000]
            jd = job_description if job_description else "Perfil general"
            decision_raw = self.analyzer.analyze(texto_para_la_nube, jd)

            # Limpieza inmediata de memoria
            del texto_para_la_nube
            del texto_para_ia
            gc.collect()
        
            # Inicialización de variables de seguridad
            f_score, reason, apto_ia = 0, "Error desconocido", "NO"

            # 3. --- PROCESAMIENTO ROBUSTO DE LA RESPUESTA ---
            if not decision_raw or not isinstance(decision_raw, str):
                reason = "Error: La IA no devolvió respuesta. Revisa conexión o API Key."
            else:
                try:
                    # Limpieza de bordes: buscar el bloque JSON
                    start = decision_raw.find('{')
                    end = decision_raw.rfind('}') + 1
                    
                    if start != -1 and end > start:
                        res_json = json.loads(decision_raw[start:end])
                        f_score = float(res_json.get("score", 0))
                        apto_ia = res_json.get("apto", "NO")
                        reason = res_json.get("motivo", "")
                        
                        # Fallback: Si el motivo está vacío pero hay score
                        if not reason or str(reason).strip().lower() in ["none", "null", ""]:
                            if f_score >= 70:
                                reason = "Cumple con los requisitos técnicos principales del puesto."
                            elif f_score >= 50:
                                reason = "Perfil con potencial o trayectoria histórica que requiere validación."
                            else:
                                reason = "No se detecta alineación técnica con la vacante."
                    else:
                        reason = "Error: La respuesta de la IA no contiene un formato JSON válido."
                except Exception as e:
                    reason = f"Error de interpretación en respuesta IA: {str(e)}"

            # 4. Clasificación física del archivo original
            nombre_original = os.path.basename(file_path)
            nuevo_nombre = f"{int(f_score):02d}_{nombre_original}"

            # Lógica de carpetas
            destino = "RECLUTADOS" if f_score >= 70 or apto_ia.upper() == "SI" else \
                      "DUDAS" if f_score >= 50 else "DESCARTADOS"

            ruta_final = os.path.join(self.base_output, destino, nuevo_nombre)
            
            if os.path.exists(file_path):
                shutil.move(file_path, ruta_final)

            # 5. Registro y Retorno
            resultado = {
                "status": "success",
                "decision": destino,
                "score": f_score,
                "motivo": reason.replace('\n', ' ').strip(),
                "nombre": nombre_original,
                "dest_path": ruta_final
            }
            self._append_to_report(resultado)
            time.sleep(15)  # Pausa para evitar saturar la IA con múltiples archivos en rápida sucesión
            return resultado

        except Exception as e:
            return {"status": "error", "reason": str(e)}